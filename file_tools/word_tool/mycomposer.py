import os
from docx import Document
from docxcompose.composer import Composer
from docxcompose.properties import CustomProperties
from copy import deepcopy
from docx.oxml.section import CT_SectPr
# from ht_script import ht

work_path = os.getcwd() + "/20200416134916/"


class MyComposer(Composer):
    def __init__(self, doc):
        super().__init__(doc)

    def replace(self, doc, remove_property_fields=True):
        """Append the given document."""
        index = self.append_index()
        self.replace_insert(
            index, doc, remove_property_fields=remove_property_fields)

    def replace_insert(self, index, doc, remove_property_fields=True):
        """Insert the given document at the given index."""
        self.reset_reference_mapping()

        # Remove custom property fields but keep the values
        if remove_property_fields:
            cprops = CustomProperties(doc)
            for name in cprops.keys():
                cprops.dissolve_fields(name)

        self._create_style_id_mapping(doc)

        for element in doc.element.body:
            if isinstance(element, CT_SectPr):
                continue
            element = deepcopy(element)
            # self.doc.element.body.insert(index, element)
            self.add_referenced_parts(doc.part, self.doc.part, element)
            self.add_styles(doc, element)
            self.add_numberings(doc, element)
            self.restart_first_numbering(doc, element)
            self.add_images(doc, element)
            self.add_shapes(doc, element)
            self.add_footnotes(doc, element)
            self.remove_header_and_footer_references(doc, element)
            index += 1

        self.add_styles_from_other_parts(doc)
        self.renumber_bookmarks()
        self.renumber_docpr_ids()
        self.fix_section_types(doc)


# def replacebookmark(src_file, dest_file, target):
#     print(dest_file)
#     composer = MyComposer(Document())
#     composer.append(Document(work_path+dest_file+".docx"))
#     composer.replace(Document(work_path+src_file+".docx"))
#     composer.save(work_path+"tmp.docx")
#     ht(src_file_name=src_file, dest_file_name="tmp",
#        res_file_name="result.docx", work_path=work_path)


if __name__ == "__main__":

    print(work_path)
    dest_file = "机器预填【部分】"
    src_file = "公司股份结构"
    # replacebookmark(src_file=src_file, dest_file=dest_file,
    #                 target="./tmp.docx")
