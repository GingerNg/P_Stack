from docx import Document
from docx.oxml.section import CT_SectPr
from docxcompose.utils import xpath
from docxcompose.composer import Composer
from docxcompose.properties import CustomProperties
from copy import deepcopy
from docx.oxml.section import CT_SectPr
import xlrd

root_path = "/home/ginger/Projects/Learning/P_Stack/file_tools/word_tool/20200416134916"


def xlsx2docx(xlsx_path, docx_path):
    readbook = xlrd.open_workbook(xlsx_path)
    sheet = readbook.sheet_by_index(0)
    nrows = sheet.nrows
    ncols = sheet.ncols
    doc = Document()
    table = doc.add_table(rows=nrows, cols=ncols)
    table.style = "Medium Grid 1 Accent 1"
    for i in range(nrows):
        cur_row_values = sheet.row_values(i)
        cur_row_cells = table.rows[i].cells
        for j in range(ncols):
            cur_row_cells[j].text = str(cur_row_values[j])
    doc.save(docx_path)


class MyComposer(Composer):
    def __init__(self, doc):
        super().__init__(doc)

    def replace(self, doc, remove_property_fields=True):
        """Append the given document."""
        index = self.append_index()
        self.replace_insert(
            index, doc, remove_property_fields=remove_property_fields)

    def replace_insert(self, index, doc, remove_property_fields=True):
        """Insert the given document at the given index."""
        self.reset_reference_mapping()

        # Remove custom property fields but keep the values
        if remove_property_fields:
            cprops = CustomProperties(doc)
            for name in cprops.keys():
                cprops.dissolve_fields(name)

        self._create_style_id_mapping(doc)

        for element in doc.element.body:
            if isinstance(element, CT_SectPr):
                continue
            element = deepcopy(element)
            # self.doc.element.body.insert(index, element)
            self.add_referenced_parts(doc.part, self.doc.part, element)
            self.add_styles(doc, element)
            self.add_numberings(doc, element)
            self.restart_first_numbering(doc, element)
            self.add_images(doc, element)
            self.add_shapes(doc, element)
            self.add_footnotes(doc, element)
            self.remove_header_and_footer_references(doc, element)
            index += 1

        self.add_styles_from_other_parts(doc)
        self.renumber_bookmarks()
        self.renumber_docpr_ids()
        self.fix_section_types(doc)


srcs = {}


def src_func(file_name):
    srcs[file_name] = []
    src_file_path = root_path + "/"+file_name+".docx"
    # 获取src文件的内容
    src_docx_obj = Document(src_file_path)
    src_element = None
    for element in src_docx_obj.element.body:
        if isinstance(element, CT_SectPr):
            continue
        # elif src_element is None:
        #     src_element = element
        # print(element)
        else:
            srcs[file_name].append(element)
        # print(element)
    # srcs[file_name] = src_element

import time
if __name__ == "__main__":
    # src_file_path = "/home/ginger/Projects/Learning/P_Stack/file_tools/word_tool/20200416134916/表格.docx"
    # # 获取src文件的内容
    # src_docx_obj = Document(src_file_path)
    # src_element = None
    # for element in src_docx_obj.element.body:
    #     if isinstance(element, CT_SectPr):
    #         continue
    #     elif src_element is None:
    #         src_element = element
    start = time.time()
    key1 = "公司股本结构"
    key2 = "公司基本资料"
    key3 = "公司合并利润表主要数据"
    key4 = "公司报告期主要财务指标"
    key5 = "公有云服务介绍"

    xlsx2docx(xlsx_path=root_path+"/"+key1+".xlsx",
              docx_path=root_path+"/"+key1+".docx")
    xlsx2docx(xlsx_path=root_path+"/"+key2+".xlsx",
              docx_path=root_path+"/"+key2+".docx")
    xlsx2docx(xlsx_path=root_path+"/"+key3+".xlsx",
              docx_path=root_path+"/"+key3+".docx")
    xlsx2docx(xlsx_path=root_path+"/"+key4+".xlsx",
              docx_path=root_path+"/"+key4+".docx")

    src_func(key1)
    src_func(key2)
    src_func(key3)
    src_func(key4)
    src_func(key5)

    composer = MyComposer(Document())
    # composer.append(Document(root_path+""+".docx"))
    composer.replace(Document(root_path+"/公有云服务介绍"+".docx"))
    composer.replace(Document(root_path+"/机器预填【部分】"+".docx"))

    dest_file_path = root_path + "/机器预填【部分】.docx"
    docx_obj = Document(dest_file_path)
    gindex = 1
    for element in docx_obj.element.body:
        if isinstance(element, CT_SectPr):
            continue
        else:
            # print(type(element))
            bookmarks_start = xpath(element, './/w:sdt')
            if bookmarks_start:
                children = element.getchildren()
                for child in children:
                    if child.tag == "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}alias":
                        # index = int(bookmarks_start[0].get(
                        #     "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"))
                        key = bookmarks_start[0].get(
                            "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val")
                        print(key)
                        src_ee = srcs.get(key)
                        if src_ee is not None:
                            element.remove(child)
                            # bookmarks_start = [src_element]
                            src_ees = srcs.get(key)
                            for src_ee in src_ees:
                                element.append(src_ee)
                                gindex += 1
                composer.doc.element.body.insert(gindex, element)
            else:
                composer.doc.element.body.insert(gindex, element)
        gindex += 1
    # composer.add_styles_from_other_parts(doc)
    # composer.renumber_bookmarks()
    # composer.renumber_docpr_ids()
    # composer.fix_section_types(doc)
    composer.doc.save(root_path+"/resssssss.docx")
    print(time.time()-start)

    # for attr, val in element.items():
    # print(attr, val)
    # bookmarks_start = xpath(element, './/w:bookmarkStart')

    # if bookmarks_start:
    #     index = int(bookmarks_start[0].get(
    #         "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}id"))
    #     key = bookmarks_start[0].get(
    #         "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}name")
    #     print(key)
    #     src_ee = srcs.get(key)
    #     if src_ee is not None:
    #         # bookmarks_start = [src_element]
    #         src_ees = srcs.get(key)
    #         # for src_ee in src_ees:
    #         docx_obj.element.body.insert(index, src_ee)

    #
